'''
class for generating a report given a guide_manager instance, info path, and template path
'''
import docxtpl
import pandas as pd
from docx.shared import Inches

import guide_manager


class report_manager:
    def __init__(self, guide_path, info_path, template_path, output_path):
        self.gm = guide_manager.guide_manager()
        self.gm.read_csv(guide_path)
        self.gm.set_xy()
        self.gm.gen_plots()
        self.d_info = {}
        self.read_info(info_path)
        self.tpl = docxtpl.DocxTemplate(template_path)
        self.generate_report(output_path)

    def read_info(self, info_path):
        '''

        :param info_path:
        :return:
        '''
        self.df_info = pd.read_csv(info_path)
        self.d_info = self.df_info.set_index('key').iloc[:, :1].T.to_dict('list')
        for k in self.d_info.iterkeys():
            self.d_info[k] = self.d_info[k][0]

    def generate_report(self, output_path):
        '''
        Generates a Report file
        :return:
        '''
        dut_start_index = self.df_info['key'].tolist().index('<DUT>')
        dut_end_index = self.df_info['key'].tolist().index('</DUT>')
        port_start_index = self.df_info['key'].tolist().index('<PORT>')
        port_end_index = self.df_info['key'].tolist().index('</PORT>')
        dut_table = self.df_info.fillna('').iloc[dut_start_index+1:dut_end_index].as_matrix()
        port_table = self.df_info.fillna('').iloc[port_start_index+1:port_end_index].as_matrix()

        sd1 = self.tpl.new_subdoc()
        sd2 = self.tpl.new_subdoc()

        table1 = sd1.add_table(rows=len(dut_table), cols=len(dut_table[0]))
        table1.style = 'TableGrid'
        for row in range(len(dut_table)):
            for col in range(len(dut_table[0])):
                table1.cell(row, col).text = dut_table[row][col]

        table2 = sd2.add_table(rows=len(port_table), cols=len(port_table[0]))
        table2.style = 'TableGrid'
        for row in range(len(port_table)):
            for col in range(len(port_table[0])):
                table2.cell(row, col).text = port_table[row][col]

        sd3 = self.tpl.new_subdoc()

        frequencies = set([])
        f_count = 0

        for p in self.gm.plot_list():
            plot_path = ''
            for interaction in self.gm.interactions:
                if interaction['plot'] == p:
                    plot_path = interaction['plot_path']
                    x_set = set(interaction['x'])
                    frequencies.update(x_set)
                    f_count = max(f_count, len(x_set))

            sd3.add_picture(plot_path, width=Inches(6.5))

        self.d_info['DUTIdentification'] = sd1
        self.d_info['PortConfiguration'] = sd2
        self.d_info['TestSummary'] = sd3

        start_frequency = min(frequencies)
        stop_frequency = max(frequencies)
        step_size = (stop_frequency-start_frequency) / float(f_count)

        self.d_info['StartFrequency'] = str(round(start_frequency*1000))+"MHz"
        self.d_info['StopFrequency'] = str(round(stop_frequency))+"GHz"
        self.d_info['StepFrequency'] = str(round(step_size*1000))+"MHz"

        self.tpl.render(self.d_info)
        self.tpl.save(output_path)